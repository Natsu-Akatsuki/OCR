# encoding:utf-8
from PyQt5.QtCore import QObject
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt
from easydict import EasyDict
from datetime import datetime
from docx.shared import Length
from pathlib import Path


class DocDump(QObject):
    """
    有关字体设置，可参考 https://zhuanlan.zhihu.com/p/61340025
    """

    def __init__(self):
        super(DocDump, self).__init__()
        self.cur_time = None
        self.template_path = str(Path().resolve() / "template")
        self.output_dir = Path().resolve() / "output" / datetime.now().strftime('%m-%d')
        self.output_dir.mkdir(parents=True, exist_ok=True)
        self.output_dir = str(self.output_dir)

    def dump_in_docs(self):

        template_in_docs = Document(self.template_path + "/" + "入库单模板.docx")
        for i in range(len(template_in_docs.paragraphs)):
            template_in_docs.paragraphs[i].line_spacing_rule = WD_LINE_SPACING.SINGLE

        table = template_in_docs.tables[0]
        table.cell(0, 6).text = self.check_id
        table.cell(0, 10).text = self.bug_date
        table.cell(1, 6).text = self.last_total_sum_money
        table.cell(1, 2).text = self.fund

        # 名称第一行
        head_item_row = 4
        for item_id in range(self.asset_count):
            table.cell(head_item_row + item_id, 1).text = self.asset_content[item_id][0]  # 固资名称
            table.cell(head_item_row + item_id, 2).text = self.asset_content[item_id][1]  # 固资单位
            table.cell(head_item_row + item_id, 3).text = str(self.asset_content[item_id][2])  # 固资单价（含税）
            table.cell(head_item_row + item_id, 4).text = self.asset_content[item_id][3]  # 固资个数
            table.cell(head_item_row + item_id, 6).text = str(self.asset_content[item_id][4])  # 固资总价格
            table.cell(head_item_row + item_id, 7).text = "11"  # 耐久/月

        table.cell(14, 6).text = self.last_total_sum_money

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.name = u'宋体'
                font.size = Pt(10)

        self.cur_time = datetime.now().strftime('%H-%M-%S')
        template_in_docs.save(self.output_dir + "/" + f'{self.cur_time}-入库单.docx')

    def dump_out_docs(self):
        template_in_docs = Document(self.template_path + "/" + "出库单模板.docx")
        for i in range(len(template_in_docs.paragraphs)):
            template_in_docs.paragraphs[i].line_spacing_rule = WD_LINE_SPACING.SINGLE
        table = template_in_docs.tables[0]
        table.cell(0, 6).text = self.bug_date

        # 名称第一行
        head_item_row = 3
        for item_id in range(self.asset_count):
            table.cell(head_item_row + item_id, 1).text = self.asset_content[item_id][0]  # 固资名称
            table.cell(head_item_row + item_id, 2).text = self.asset_content[item_id][1]  # 固资单位
            table.cell(head_item_row + item_id, 3).text = str(self.asset_content[item_id][2])  # 固资单价（含税）
            table.cell(head_item_row + item_id, 4).text = self.asset_content[item_id][3]  # 固资个数
            table.cell(head_item_row + item_id, 5).text = str(self.asset_content[item_id][4])  # 固资总价格

        table.cell(15, 4).text = self.last_total_sum_money

        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                paragraph = paragraphs[0]
                run_obj = paragraph.runs
                run = run_obj[0]
                font = run.font
                font.name = u'宋体'
                font.size = Pt(10)

        template_in_docs.save(self.output_dir + "/" + f'{self.cur_time}-出库单.docx')

    def run(self, ocr_data, fund_id):

        '''
        :param dict ocr_data:
             "InvoiceNum"
             "InvoiceDate"
             "obj_info"
        '''
        ocr_data = EasyDict(ocr_data)
        self.check_id = ocr_data.invoice_nums  # 发票号码
        self.bug_date = ocr_data.invoice_dates.replace("年", '.').replace("月", '.').replace("日", '')  # 购置日期
        self.fund = fund_id
        self.asset_count = len(ocr_data.obj_infos)  # 固资类数
        self.asset_content = ocr_data.obj_infos  # 固资信息
        self.last_total_sum_money = ocr_data.amount_in_figuers  # 申报固资总价钱
        self.dump_in_docs()
        self.dump_out_docs()


if __name__ == '__main__':
    test_dict = {'invoice_nums': '35782608\n52366743', 'invoice_dates': '2020年11月21日\n2017年03月23日',
                 'amount_in_figuers': '759.6',
                 'obj_infos': [('*计算机配套产品*电源', '台', 259.0, '1', 259.0), ('卡托普利片(硫甲丙脯酸片)', '瓶', 2.6, '10', 26.0),
                               ('咳特灵胶囊', '瓶', 4.5, '30', 135.0), ('皮肤病血毒丸', '瓶', 8.1, '10', 81.0),
                               ('酒石酸美托洛尔片(曲新克治)', '盒', 3.6, '20', 72.0), ('小儿七星茶颗粒', '盒', 0, '0', 66.0),
                               ('单硝酸异山梨酯片', '盒', 7.5, '10', 75.0), ('防风通圣丸', '盒', 2.95, '8', 23.6),
                               ('维生素B1片', '瓶', 2.2, '10', 22.0)]}

    doc_dump = DocDump()
    doc_dump.run(test_dict, "0090")
